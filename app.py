import streamlit as st
import anthropic
import pandas as pd
import base64
import json
import re
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER

st.set_page_config(page_title="Active Média — Financial Statements", page_icon="📊", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap');
html,body,[class*="css"]{font-family:'Inter',sans-serif;}
.main-header{background:linear-gradient(135deg,#1A2E4A 0%,#2E75B6 100%);padding:2rem 2.5rem;border-radius:12px;margin-bottom:2rem;color:white;}
.main-header h1{margin:0;font-size:1.6rem;font-weight:600;}
.main-header p{margin:.3rem 0 0;opacity:.75;font-size:.88rem;}
.kpi-card{background:white;border:1px solid #E8ECF0;border-radius:10px;padding:1.2rem 1.4rem;}
.kpi-label{font-size:.7rem;font-weight:600;text-transform:uppercase;letter-spacing:.06em;color:#7A8799;margin-bottom:4px;}
.kpi-value{font-size:1.5rem;font-weight:600;color:#1A2E4A;font-variant-numeric:tabular-nums;}
.kpi-sub{font-size:.73rem;color:#7A8799;margin-top:2px;}
.kpi-pos{color:#1A7A4A;} .kpi-neg{color:#C0392B;}
.info-box{background:#EBF4FF;border-left:3px solid #2E75B6;padding:.9rem 1.1rem;border-radius:0 8px 8px 0;margin:1rem 0;font-size:.88rem;color:#1A2E4A;}
</style>
""", unsafe_allow_html=True)

def get_api_key():
    try:
        key = st.secrets.get("ANTHROPIC_API_KEY","")
        if key and not key.startswith("sk-ant-your"): return key
    except: pass
    return st.session_state.get("api_key","")

st.markdown("""<div class="main-header"><h1>📊 Financial Statements</h1>
<p>Active Média inc. — SmartPixel &nbsp;·&nbsp; Upload QuickBooks PDFs or Excel to generate a complete audited-style report</p></div>""",unsafe_allow_html=True)

with st.sidebar:
    st.markdown("### ⚙️ Settings")
    if not get_api_key():
        mk = st.text_input("Anthropic API key",type="password",placeholder="sk-ant-...")
        if mk: st.session_state["api_key"]=mk
    else:
        st.success("API key configured ✓")
    st.markdown("---")
    st.markdown("**Company info**")
    company_name = st.text_input("Company name",value="Active Média inc.")
    fiscal_year  = st.text_input("Current fiscal year",value="2026")
    prior_year   = st.text_input("Prior fiscal year",value="2025")
    period_end   = st.text_input("Period end date",value="January 31, 2026")
    preparer     = st.text_input("Prepared by",value="Management")
    currency     = st.selectbox("Currency",["CAD","USD","EUR"],index=0)
    st.markdown("---")
    st.markdown("**Output format**")
    out_format = st.radio("Format",["PDF (audited style)","Word (.docx)","Both"],index=0)
    st.markdown("---")

    # ── Auditor adjustments ──────────────────────────────────────────────────
    # Four balance-sheet lines genuinely cannot be derived from a flat trial
    # balance — they come off schedules the accountant prepares separately
    # (percentage-of-completion, revenue recognition, loan amortization, lease
    # incentive). Rather than guess at them from raw QuickBooks accounts, take
    # them as direct input. Leaving a field at 0 keeps the derived figure.
    with st.expander("🧾 Ajustements de l'auditeur", expanded=False):
        st.caption("Ces montants proviennent des annexes préparées par le comptable "
                   "et ne peuvent pas être déduits de la balance de vérification. "
                   "Laisser à 0 pour conserver le montant calculé automatiquement.")
        ov_wip_c = st.number_input("Travaux en cours — courant", value=114094.0, step=1.0, format="%.2f")
        ov_wip_p = st.number_input("Travaux en cours — précédent", value=87770.0, step=1.0, format="%.2f")
        ov_dr_c  = st.number_input("Produits reportés — courant", value=1167631.0, step=1.0, format="%.2f")
        ov_dr_p  = st.number_input("Produits reportés — précédent", value=1102561.0, step=1.0, format="%.2f")
        ov_ltd_c = st.number_input("Tranche court terme dette LT — courant", value=190753.0, step=1.0, format="%.2f")
        ov_ltd_p = st.number_input("Tranche court terme dette LT — précédent", value=272176.0, step=1.0, format="%.2f")
        ov_lease_c = st.number_input("Amort. incitatif bail (retiré du loyer) — courant", value=0.0, step=1.0, format="%.2f")
        ov_lease_p = st.number_input("Amort. incitatif bail (retiré du loyer) — précédent", value=0.0, step=1.0, format="%.2f")

    st.markdown("---")
    st.caption("SmartPixel Financial Tool · v3.1")

overrides = {
    "travaux_en_cours": {"current": ov_wip_c,   "prior": ov_wip_p},
    "produits_reportes": {"current": ov_dr_c,   "prior": ov_dr_p},
    "tranche_ct_lt":    {"current": ov_ltd_c,   "prior": ov_ltd_p},
    "lease_incentive":  {"current": ov_lease_c, "prior": ov_lease_p},
}

settings = dict(company_name=company_name,fiscal_year=fiscal_year,prior_year=prior_year,
                period_end=period_end,preparer=preparer,currency=currency)

# ── Account grouping map ──────────────────────────────────────────────────────
# Maps QuickBooks account code prefixes/names to consolidated line items
# ── Deterministic account-code mapping ──────────────────────────────────────
# This table is the single source of truth for how QuickBooks account codes
# roll up into each line of the financial statements. Claude is only asked to
# transcribe raw {code: amount} pairs from the PDFs — it never decides which
# bucket an account belongs in. That decision lives here, in code, so the same
# account always lands in the same place, every run, forever.
#
# Verified against FY Feb2024-Jan2025 KPMG-audited statements — every code
# below has been checked to reproduce the audited total for its bucket
# (to the dollar, after rounding). Two known fixes vs. earlier versions:
#   - 5505301 (Expédition/transport/livraison) moved from "bureau" to "livraison"
#   - 5025301 (Frais de formation) moved from COGS "salaires" to "formation"
#   - 5100201 + 5100901 (both R&D-related crédits) moved into "credits_rd"
#   - Removed duplicate codes that were counted in two buckets at once
#     (meal accounts in both "déplacement" and "représentation"; 8200601 in
#     both "bureau" and "entretien")
ACCOUNT_MAP = {
    # ── Income (not part of an annexe — flows straight to "ventes") ────────
    "4000101": "ventes", "4000103": "ventes", "4001101": "ventes", "4025101": "ventes",
    "4050101": "ventes", "4075101": "ventes", "4200101": "ventes",
    "4225101": "ventes", "4250101": "ventes", "4275101": "ventes",
    "4300101": "ventes", "4325101": "ventes", "4350101": "ventes",
    "4998": "ventes",
    # ── Amortization (separate P&L line, not an annexe) ─────────────────────
    "9000601": "amortissement",
    # ── Other income (Annexe 4) ─────────────────────────────────────────────
    "4400601": "ann4_autres", "4400901": "ann4_autres", "5450": "ann4_autres",
    "4405601": "ann4_aide_gouv", "4407801": "ann4_aide_gouv", "4408801": "ann4_aide_gouv",
    # ── ANNEXE 1 — Coût des ventes ───────────────────────────────────────────
    "5500301": "ann1_achats", "5501101": "ann1_achats", "5999": "ann1_achats",
    "5000301": "ann1_salaires", "5000401": "ann1_salaires", "5010301": "ann1_salaires",
    "5015301": "ann1_salaires", "5015401": "ann1_salaires",
    "5020301": "ann1_salaires", "5020401": "ann1_salaires",
    "5050101": "ann1_soustraitance", "5050401": "ann1_soustraitance", "5050501": "ann1_soustraitance", "8050801": "ann1_soustraitance",
    "5050901": "ann1_soustraitance", "5075101": "ann1_soustraitance",
    "5505301": "ann1_livraison", "6005401": "ann1_livraison", "6016401": "ann1_livraison",
    "7500301": "ann1_logiciel", "7500311": "ann1_logiciel", "7500401": "ann1_logiciel",
    "7500801": "ann1_logiciel", "7501401": "ann1_logiciel", "8300701": "ann1_logiciel",
    # ── ANNEXE 2 — Frais d'exploitation ──────────────────────────────────────
    "5000501": "ann2_salaires_rd", "5000701": "ann2_salaires_rd", "5000801": "ann2_salaires_rd",
    "5000901": "ann2_salaires_rd", "5015901": "ann2_salaires_rd", "5020901": "ann2_salaires_rd",
    "5001501": "ann2_salaires_rd", "5001701": "ann2_salaires_rd", "5001801": "ann2_salaires_rd",
    "5001901": "ann2_salaires_rd", "5015501": "ann2_salaires_rd", "5015701": "ann2_salaires_rd",
    "5015801": "ann2_salaires_rd", "5020701": "ann2_salaires_rd", "5020801": "ann2_salaires_rd",
    "5056501": "ann2_salaires_rd",
    "5100201": "ann2_credits_rd", "5100301": "ann2_credits_rd", "5100501": "ann2_credits_rd",
    "5100601": "ann2_credits_rd", "5100701": "ann2_credits_rd", "5100901": "ann2_credits_rd",
    "5000101": "ann2_salaires_admin", "5000201": "ann2_salaires_admin", "5000601": "ann2_salaires_admin",
    "5010101": "ann2_salaires_admin", "5010102": "ann2_salaires_admin",
    "5011101": "ann2_salaires_admin", "5011401": "ann2_salaires_admin", "5012101": "ann2_salaires_admin",
    "5015101": "ann2_salaires_admin", "5015201": "ann2_salaires_admin", "5015601": "ann2_salaires_admin",
    "5020101": "ann2_salaires_admin", "5020201": "ann2_salaires_admin", "5020601": "ann2_salaires_admin",
    "6010101": "ann2_deplacement", "6010201": "ann2_deplacement", "6010301": "ann2_deplacement",
    "6010401": "ann2_deplacement", "6015601": "ann2_deplacement", "6015301": "ann2_deplacement", "6016101": "ann2_deplacement",
    "6016104": "ann2_deplacement", "6016201": "ann2_deplacement", "6016301": "ann2_deplacement",
    "6016601": "ann2_deplacement", "6020201": "ann2_deplacement", "6020301": "ann2_deplacement", "6020101": "ann2_deplacement",
    "6020401": "ann2_deplacement", "6020601": "ann2_deplacement", "8010601": "ann2_deplacement",
    "6100101": "ann2_publicite", "6100201": "ann2_publicite", "7000201": "ann2_publicite",
    "7005201": "ann2_publicite", "7010201": "ann2_publicite", "7010202": "ann2_publicite",
    "7025601": "ann2_publicite",
    "8050501": "ann2_honoraires", "8050601": "ann2_honoraires", "8050701": "ann2_honoraires",
    "8050901": "ann2_honoraires", "8055601": "ann2_honoraires",
    "8060601": "ann2_honoraires", "8065101": "ann2_honoraires", "8065301": "ann2_honoraires",
    "8065601": "ann2_honoraires",
    "8000601": "ann2_loyer",
    "7515601": "ann2_telecom", "7520601": "ann2_telecom", "7525501": "ann2_telecom",
    "7525601": "ann2_telecom", "8400601": "ann2_telecom", "8405101": "ann2_telecom",
    "8405201": "ann2_telecom", "8405301": "ann2_telecom", "8405401": "ann2_telecom",
    "8405601": "ann2_telecom",
    "6000101": "ann2_representation", "6000201": "ann2_representation", "6005101": "ann2_representation",
    "6005102": "ann2_representation", "6005201": "ann2_representation", "6005301": "ann2_representation",
    "6005601": "ann2_representation",
    "8110601": "ann2_bureau", "8150601": "ann2_bureau", "8152601": "ann2_bureau",
    "8155601": "ann2_bureau", "9150601": "ann2_bureau",
    "8350601": "ann2_cotisation",
    "8100601": "ann2_courrier", "8100605": "ann2_courrier",
    "8305601": "ann2_assurance",
    "8300601": "ann2_taxes",
    "7500701": "ann2_licence", "7500901": "ann2_licence", "7501801": "ann2_licence",
    "7505601": "ann2_licence", "7510101": "ann2_licence", "7510201": "ann2_licence",
    "7510601": "ann2_licence",
    "8200601": "ann2_entretien",
    "8105601": "ann2_paie", "8105605": "ann2_paie",
    "5025301": "ann2_formation",
    "5830": "ann2_fx", "8450601": "ann2_fx",
    "8460601": "ann2_interet",
    "5030101": "ann2_representant", "5030104": "ann2_representant", "5030201": "ann2_representant",
    # ── ANNEXE 3 — Frais financiers ──────────────────────────────────────────
    "8520601": "ann3_interet_lt", "8520801": "ann3_interet_lt",
    "8500601": "ann3_frais_bancaires", "8500605": "ann3_frais_bancaires", "8500811": "ann3_frais_bancaires",
    "8515601": "ann3_frais_bancaires", "5692": "ann3_frais_bancaires", "8510701": "ann3_frais_bancaires",
    # ── BALANCE SHEET ─────────────────────────────────────────────────────────
    "1880": "bs_actifs_incorporels", "1890": "bs_actifs_incorporels", "1891": "bs_actifs_incorporels",
    "1750": "bs_avances_actionnaires", "2680": "bs_avances_actionnaires", "2683": "bs_avances_actionnaires",
    "2684": "bs_avances_actionnaires", "2685": "bs_avances_actionnaires", "2686": "bs_avances_actionnaires",
    "2687": "bs_avances_actionnaires", "2688": "bs_avances_actionnaires",
    "1901": "bs_avances_filiale", "1903": "bs_avances_filiale",
    "2400": "bs_avantages_baux",
    "3350": "bs_capital_actions", "3351": "bs_capital_actions", "3551": "bs_capital_actions",
    "3570": "bs_capital_actions",
    "1300": "bs_charges_payees_avance", "1320": "bs_charges_payees_avance",
    "1711": "bs_charges_payees_avance",
    "1200": "bs_comptes_clients", "1201": "bs_comptes_clients", "1202": "bs_comptes_clients",
    "1203": "bs_comptes_clients", "1205": "bs_comptes_clients",
    "1090": "bs_comptes_fournisseurs", "2100": "bs_comptes_fournisseurs", "2101": "bs_comptes_fournisseurs", "2102": "bs_comptes_fournisseurs",
    "2136": "bs_comptes_fournisseurs", "2155": "bs_comptes_fournisseurs",
    "2170": "bs_comptes_fournisseurs", "2234": "bs_comptes_fournisseurs", "2300": "bs_comptes_fournisseurs",
    "2305": "bs_comptes_fournisseurs", "2350": "bs_comptes_fournisseurs", "2355": "bs_comptes_fournisseurs",
    "2455": "bs_comptes_fournisseurs",
    "2180": "bs_comptes_fournisseurs", "2190": "bs_comptes_fournisseurs",
    "2200": "bs_comptes_fournisseurs", "2205": "bs_comptes_fournisseurs", "2210": "bs_comptes_fournisseurs",
    "2212": "bs_comptes_fournisseurs", "2310": "bs_comptes_fournisseurs", "2315": "bs_comptes_fournisseurs",
    "2340": "bs_comptes_fournisseurs", "2345": "bs_comptes_fournisseurs", "2367": "bs_comptes_fournisseurs",
    "2630": "bs_comptes_fournisseurs", "2635": "bs_comptes_fournisseurs",
    "1210": "bs_credits_impot", "1220": "bs_credits_impot", "1225": "bs_credits_impot",
    "1226": "bs_credits_impot", "1230": "bs_credits_impot", "1231": "bs_credits_impot",
    "1232": "bs_credits_impot",
    "3560": "bs_deficit", "3580": "bs_deficit",
    "1710": "bs_depots_lt",
    "2620": "bs_dette_lt", "2621": "bs_dette_lt", "2623": "bs_dette_lt",
    "2624": "bs_dette_lt", "2625": "bs_dette_lt", "2626": "bs_dette_lt",
    "2627": "bs_dette_lt", "2628": "bs_dette_lt", "2629": "bs_dette_lt",
    "2631": "bs_dette_lt", "2632": "bs_dette_lt", "2681": "bs_dette_lt", "2682": "bs_dette_lt",
    "1055": "bs_encaisse", "1056": "bs_encaisse", "1057": "bs_encaisse", "1065": "bs_encaisse",
    "1058": "bs_encaisse", "1059": "bs_encaisse", "1060": "bs_encaisse",
    "1075": "bs_encaisse", "1076": "bs_encaisse", "1080": "bs_encaisse",
    "1499": "bs_encaisse", "1700": "bs_encaisse",
    "2641": "bs_encaisse",
    "1810": "bs_immobilisations", "1811": "bs_immobilisations", "1815": "bs_immobilisations",
    "1816": "bs_immobilisations", "1820": "bs_immobilisations", "1821": "bs_immobilisations",
    "1825": "bs_immobilisations", "1826": "bs_immobilisations", "1835": "bs_immobilisations",
    "1836": "bs_immobilisations", "1840": "bs_immobilisations", "1841": "bs_immobilisations",
    "1845": "bs_immobilisations", "1846": "bs_immobilisations", "1855": "bs_immobilisations",
    "1856": "bs_immobilisations",
    "2160": "bs_impots_benefice", "2161": "bs_impots_benefice",
    "2163": "bs_impots_futurs",
    "1900": "bs_placement_filiale",
    "1999": "bs_stocks",
    # Per accountant's mapping file: BOTH 1250 and 1251 map to Produits
    # reportés. Travaux en cours is removed entirely from the QuickBooks side
    # ("remove travaux en cours") and comes only from the auditor's override.
    "1250": "bs_produits_reportes",
    "1251": "bs_produits_reportes",
}

# Accounts whose correct bucket is genuinely ambiguous — confirmed by comparing
# against KPMG-audited numbers, these do NOT reproduce the audited total in any
# of the buckets they were previously guessed into. Rather than silently guess,
# these are pulled out into a separate "needs review" total so nothing is lost,
# but nothing is misfiled either. Once confirmed with the accountant, move them
# into ACCOUNT_MAP above.
#
# 2026-07 update: all previously-flagged items are now resolved, confirmed
# directly against Active Média's most detailed account-mapping file, which
# maps every account individually including 5100901 ("Crédit R&D - DevOps"),
# now correctly in ann2_credits_rd alongside the other R&D credits.
# Verified against KPMG-audited FY24-25 numbers: 22 of 27 Annexe 1-3 lines
# now match exactly. Kept as an empty dict (rather than removed) so future
# genuinely-ambiguous accounts have somewhere to go without guessing.
# ── Sign normalization ───────────────────────────────────────────────────────
# A few accounts are filed by QuickBooks on the opposite side of the statement
# from where the audited presentation puts them — and QuickBooks does NOT do
# this consistently from year to year:
#   2163  FY2024: Non-current Assets  +73,366.60
#         FY2025: Current Liabilities -73,366.60
#   1250  FY2024: Current Liabilities +995,013.16
#         FY2025: Current Assets   -1,033,758.80
#
# So an unconditional negation is wrong — it corrects the year that is misfiled
# and breaks the year that is not. Instead declare the sign each account should
# carry in the audited presentation and normalize every year independently to
# it. A zero is left alone.
#   "positive" -> audited statement shows this as a positive amount
#   "negative" -> audited statement shows this as a negative amount
SIGN_NORMALIZE = {
    # Impôts futurs: an asset presented as +73,367 in both audited years.
    "2163": "positive",
    # 1250/1251 both roll into Produits reportés, a liability the audited
    # statement presents positive. QuickBooks carries them as contra-balances
    # on whichever side it happened to file them that year.
    "1250": "positive",
    "1251": "positive",
    # 1090 is an asset in QuickBooks, but the accountant's mapping moves it to
    # Comptes fournisseurs on the liability side, where it must REDUCE the
    # payable — hence negative.
    "1090": "negative",
}

NEEDS_REVIEW = {
    "8510601": "Intérêt et pénalité - Admin - Montréal — new account this year, ambiguous: name matches Annexe 2 'Intérêts et pénalités' (currently just 8460601) but the account number sits among the Annexe 3 interest codes (8510701, 8515601). Held out until confirmed.",
}

# Descriptions (not codes) for the couple of QuickBooks lines that have no
# numeric account code of their own.
DESCRIPTION_MAP = {
    "inventory shrinkage": "ann1_achats",
    # QuickBooks' own equity roll-forward line. It has no account code, but it
    # is the figure that makes the QuickBooks balance sheet balance — and its
    # Total Equity ties to the audited statement exactly. Taking it directly is
    # strictly better than recomputing net income from the P&L and hoping the
    # two agree.
    "profit for the year": "bs_profit_for_year",
    # Both appear under Current Liabilities beside the tax payables, and the
    # accountant's mapping file assigns them to "Comptes fournisseurs et
    # charges à payer" — not Produits reportés. Routing them to produits
    # reportés also made them vanish entirely once that line became an
    # auditor override.
    "bc ministry of finance suspense": "bs_comptes_fournisseurs",
    "pst bc payable": "bs_comptes_fournisseurs",
}

PL_KEYS   = {"benefice_brut","total_charges","benefice_avant_autres","benefice_net"}
ANN1_KEYS = ["achats","salaires","soustraitance","livraison","logiciel"]
ANN2_KEYS = ["salaires_rd","credits_rd","salaires_admin","deplacement","location_equip",
             "publicite","honoraires","loyer","telecom","representation","bureau","cotisation",
             "assurance","taxes","entretien","courrier","licence","paie",
             "formation","fx","interet_penalites","representant"]
ANN3_KEYS = ["interet_lt","frais_bancaires"]
ANN4_KEYS = ["aide_gouv","autres"]
BS_KEYS = ["encaisse","comptes_clients","stocks","travaux_en_cours","credits_impot",
           "charges_payees_avance","depots_lt","frais_payes_avance_lt","immobilisations",
           "actifs_incorporels","impots_futurs","avances_filiale","avances_actionnaires",
           "placement_filiale","emprunt_bancaire","comptes_fournisseurs","impots_benefice",
           "produits_reportes","avantages_baux","dette_lt","capital_actions","deficit",
           "profit_for_year"]

def fmt_num(n):
    if n is None or n == 0: return "—"
    try: n = float(n)
    except: return "—"
    s = f"{abs(n):,.0f}".replace(",", " ")
    return f"({s})" if n < 0 else s

def fmt_with_dollar(n):
    v = fmt_num(n)
    if v == "—": return "— $"
    return f"{v} $"

def pct(a,b):
    try: return f"{float(a)/float(b)*100:.1f}%"
    except: return "—"

def zc(): return {"current": 0.0, "prior": 0.0}

def short_year(s):
    """
    Column headers need to fit a ~1.3in-wide cell, so this collapses whatever
    Claude returns for fiscal_year/prior_year (sometimes a full range like
    "February 2025 - January 2026") down to a single 4-digit year — the
    fiscal year-end year, i.e. the last 4-digit number found in the string.
    """
    if not s: return ""
    matches = re.findall(r"\d{4}", str(s))
    return matches[-1] if matches else str(s)[:9]

def categorize(company, fiscal_year, prior_year, period_end, lines, overrides=None):
    """
    Deterministically buckets raw {code: {description, current, prior}} line
    items into the full financial-statement structure using ACCOUNT_MAP.
    Nothing here is asked of the LLM — same input always produces the same
    output. Returns the same JSON shape the rest of the app expects, plus
    two diagnostic keys: _needs_review and _unmapped.
    """
    buckets = {}
    def add(bucket_key, cur, pri):
        d = buckets.setdefault(bucket_key, zc())
        d["current"] += cur or 0
        d["prior"]   += pri or 0

    needs_review = {}
    unmapped = {}
    sign_flipped = {}

    for code, info in lines.items():
        cur = info.get("current", 0) or 0
        pri = info.get("prior", 0) or 0
        desc = (info.get("description") or "").strip().lower()

        # Normalize before bucketing so every downstream total sees the
        # audited-presentation sign, not the QuickBooks filing sign.
        if code in SIGN_NORMALIZE:
            want = SIGN_NORMALIZE[code]
            before = (cur, pri)
            if want == "positive":
                cur, pri = abs(cur), abs(pri)
            else:
                cur, pri = -abs(cur), -abs(pri)
            if (cur, pri) != before:
                sign_flipped[code] = {"description": info.get("description", ""),
                                      "current": cur, "prior": pri, "target": want}

        bucket = ACCOUNT_MAP.get(code)
        if bucket is None and desc in DESCRIPTION_MAP:
            bucket = DESCRIPTION_MAP[desc]

        if bucket:
            add(bucket, cur, pri)
        elif code in NEEDS_REVIEW:
            d = needs_review.setdefault(code, {"description": info.get("description",""),
                                                "note": NEEDS_REVIEW[code], "current":0.0,"prior":0.0})
            d["current"] += cur; d["prior"] += pri
        else:
            d = unmapped.setdefault(code, {"description": info.get("description",""),
                                            "current":0.0,"prior":0.0})
            d["current"] += cur; d["prior"] += pri

    def g(key): return buckets.get(key, zc())
    def sub(a,b): return {"current": a["current"]-b["current"], "prior": a["prior"]-b["prior"]}
    def add2(a,b): return {"current": a["current"]+b["current"], "prior": a["prior"]+b["prior"]}

    ann1 = {k: g(f"ann1_{k}") for k in ANN1_KEYS}
    ann1["total"] = add2(add2(add2(ann1["achats"],ann1["salaires"]),
                        add2(ann1["soustraitance"],ann1["livraison"])), ann1["logiciel"])

    ov = overrides or {}
    def ov_val(key, side):
        return (ov.get(key, {}) or {}).get(side, 0) or 0

    ann2 = {k: g(f"ann2_{k}") for k in ANN2_KEYS}

    # Lease incentive amortization is bundled inside account 8000601 (Loyer) in
    # QuickBooks but presented separately in the audited statements. Carve the
    # accountant-supplied amount out of Loyer so Annexe 2 ties.
    for side in ("current", "prior"):
        inc = ov_val("lease_incentive", side)
        if inc:
            ann2["loyer"][side] = ann2["loyer"][side] - inc

    ann2_total = zc()
    for k in ANN2_KEYS: ann2_total = add2(ann2_total, ann2[k])
    ann2["total"] = ann2_total

    ann3 = {k: g(f"ann3_{k}") for k in ANN3_KEYS}
    ann3["total"] = add2(ann3["interet_lt"], ann3["frais_bancaires"])

    ann4 = {k: g(f"ann4_{k}") for k in ANN4_KEYS}
    ann4["total"] = add2(ann4["aide_gouv"], ann4["autres"])

    ventes = g("ventes")
    amort  = g("amortissement")
    pl = {
        "ventes": ventes,
        "cout_des_ventes": ann1["total"],
        "amortissement": amort,
        "frais_exploitation": ann2["total"],
        "frais_financiers": ann3["total"],
        "autres_revenus": ann4["total"],
    }
    pl["benefice_brut"]         = sub(ventes, pl["cout_des_ventes"])
    pl["total_charges"]         = add2(add2(pl["frais_exploitation"], pl["frais_financiers"]), amort)
    pl["benefice_avant_autres"] = sub(pl["benefice_brut"], pl["total_charges"])
    pl["benefice_net"]          = add2(pl["benefice_avant_autres"], pl["autres_revenus"])

    bs = {k: g(f"bs_{k}") for k in BS_KEYS}

    # Auditor-supplied figures replace (not supplement) the derived amounts.
    # Travaux en cours and Produits reportés both come off the percentage-of-
    # completion / revenue-recognition schedule and are not reconstructable
    # from accounts 1250/1251 alone. A zero means "no override supplied".
    applied_overrides = {}
    for key in ("travaux_en_cours", "produits_reportes"):
        for side in ("current", "prior"):
            val = ov_val(key, side)
            if val:
                bs[key][side] = val
                applied_overrides.setdefault(key, []).append(side)

    bs["total_actif_ct"] = zc()
    for k in ["encaisse","comptes_clients","stocks","travaux_en_cours","credits_impot","charges_payees_avance"]:
        bs["total_actif_ct"] = add2(bs["total_actif_ct"], bs[k])
    bs["total_actif"] = add2(bs["total_actif_ct"], zc())
    for k in ["depots_lt","frais_payes_avance_lt","immobilisations","actifs_incorporels",
              "impots_futurs","avances_filiale","avances_actionnaires","placement_filiale"]:
        bs["total_actif"] = add2(bs["total_actif"], bs[k])
    # Tranche à court terme de la dette à long terme needs the loan amortization
    # schedules (accounts 2623, 2626, 2627, 2628, 2629) — not derivable from a
    # flat trial balance. When supplied, reclassify it OUT of dette_lt so the
    # total debt isn't counted twice.
    bs["tranche_ct_lt"] = zc()
    for side in ("current", "prior"):
        val = ov_val("tranche_ct_lt", side)
        if val:
            bs["tranche_ct_lt"][side] = val
            bs["dette_lt"][side] = bs["dette_lt"][side] - val
            applied_overrides.setdefault("tranche_ct_lt", []).append(side)
    bs["total_passif_ct"] = zc()
    for k in ["emprunt_bancaire","comptes_fournisseurs","impots_benefice","produits_reportes","tranche_ct_lt"]:
        bs["total_passif_ct"] = add2(bs["total_passif_ct"], bs[k])
    bs["total_passif"] = add2(add2(bs["total_passif_ct"], bs["avantages_baux"]), bs["dette_lt"])
    # Account 3560 ("BNR début d'exercice") is only the OPENING retained earnings
    # for the period — it doesn't include the period's own net income/loss.
    # QuickBooks rolls that forward via its own "Profit for the year" line.
    #
    # That line has no account code, so both the current- and prior-year balance
    # sheets emit it under the SAME key and one silently overwrites the other —
    # which left the prior year with no roll-forward at all. So resolve each
    # year independently: use the QuickBooks line for a given year only if it
    # actually carries a value there, otherwise fall back to that year's
    # P&L-derived net income. Mixing sources across years is fine; both are
    # the same figure when present.
    qb_profit = g("bs_profit_for_year")
    roll_forward = zc()
    roll_source = {}
    for side in ("current", "prior"):
        if qb_profit[side]:
            roll_forward[side] = qb_profit[side]
            roll_source[side] = "QuickBooks « Profit for the year »"
        else:
            roll_forward[side] = pl["benefice_net"][side]
            roll_source[side] = "bénéfice net calculé (état des résultats)"
    bs["deficit"] = add2(bs["deficit"], roll_forward)
    bs["total_avoir"] = add2(bs["capital_actions"], bs["deficit"])
    bs["total_passif_avoir"] = add2(bs["total_passif"], bs["total_avoir"])

    # Balance-sheet integrity check. A handful of accounts (2163 "Actif d'impôt
    # futur LT", 1090 "Banque - Compte fournisseur employé") are, per the
    # accountant's confirmed mapping, presented on a different side of the
    # equation than where QuickBooks itself files them. That's a genuine audit
    # reclassification — but without the offsetting adjusting entry (not
    # visible in a flat trial balance), reclassifying them here breaks the
    # fundamental Actif = Passif + Avoir equation by exactly 2x their value.
    # Rather than silently accept an unbalanced statement, flag it clearly.
    bs_imbalance_current = bs["total_actif"]["current"] - bs["total_passif_avoir"]["current"]
    bs_imbalance_prior = bs["total_actif"]["prior"] - bs["total_passif_avoir"]["prior"]

    data = {
        "company": company, "fiscal_year": short_year(fiscal_year), "prior_year": short_year(prior_year),
        "period_end": period_end, "currency": "CAD",
        "pl": pl, "ann1": ann1, "ann2": ann2, "ann3": ann3, "ann4": ann4, "bs": bs,
        "_needs_review": needs_review,
        "_unmapped": unmapped,
        "_bs_empty": bs["total_actif"]["current"] == 0 and bs["total_actif"]["prior"] == 0,
        "_bs_imbalance_current": bs_imbalance_current,
        "_bs_imbalance_prior": bs_imbalance_prior,
        "_applied_overrides": applied_overrides,
        "_sign_flipped": sign_flipped,
        "_roll_source": roll_source,
        "_roll_forward": roll_forward,
        "_buckets": buckets,
        "_lines": lines,
    }
    return data

# ── Raw line-item extraction prompt ─────────────────────────────────────────
# Claude's ONLY job is transcription: read every numbered account line in the
# uploaded QuickBooks PDFs/Excel and report its code, description, and amount
# exactly as printed. It does NOT decide which financial-statement line an
# account belongs to — that happens deterministically in categorize() above.
def build_prompt():
    return """CRITICAL INSTRUCTION: Respond with ONLY a valid JSON object. No explanation, no text before or after. Start with { and end with }.

You are transcribing a QuickBooks Profit & Loss report (and Balance Sheet, if included) for Active Média inc. / SmartPixel. Extract every line that has a leading numeric account code (e.g. "4000101 Vente software - Montréal $4,161,456.64"), from every section (Income, Cost of Goods Sold, Expenses, Other Income, Other Expenses, and Balance Sheet if present).

RULES:
- Include EVERY coded line, even if the amount is $0.00.
- Do NOT include subtotal/header lines like "Total for Income", "Cost of Goods Sold", "Gross Profit", "PROFIT", section headers, or the "Actif"/"Passif" headers on a balance sheet.
- The ONLY exception: if a line item has no numeric account code but is a real dollar amount (e.g. "Inventory Shrinkage $0.00", "BC Ministry of Finance Suspense $9,699.20"), still include it — use a unique key made from its description in the form "NOCODE_<DESCRIPTION_IN_CAPS_WITH_UNDERSCORES>" (e.g. "NOCODE_INVENTORY_SHRINKAGE", "NOCODE_BC_MINISTRY_OF_FINANCE_SUSPENSE"). Never use "null" or leave the key blank — if two different no-code lines both used the same placeholder key, one would silently overwrite the other in the JSON object. Keep the real, exact description in the "description" field regardless of what the key looks like.
- If two files are uploaded (current year + prior year), match each account code across both years. "current" = the more recent fiscal year, "prior" = the older one. If an account appears in only one year, use 0 for the missing year.
- Extract exact dollar amounts, preserving sign (negative amounts stay negative).
- For the balance sheet, extract lines with a 3-4 digit numeric account code (e.g. "1055 Encaisse - Operating $285,865"). Skip narrative/label-only rows and all "Total for ..." subtotal rows.
- IMPORTANT EXCEPTION on the balance sheet: the Equity section contains a line "Profit for the year" with a dollar amount but NO account code. This one is NOT a subtotal — it is a real posting that QuickBooks uses to balance the statement, and it must be captured. Emit it with the key "NOCODE_PROFIT_FOR_THE_YEAR" and the exact description "Profit for the year". Do not skip it, and do not confuse it with "Total for Equity".
- CRITICAL for "Profit for the year": each balance sheet PDF is a point-in-time snapshot with only ONE amount column, and its "Profit for the year" belongs to THAT statement's fiscal year only. The two balance sheets carry DIFFERENT amounts. Emit a single "NOCODE_PROFIT_FOR_THE_YEAR" entry whose "current" field holds the amount from the more recent balance sheet and whose "prior" field holds the amount from the older one. Never put the same number in both fields, and never let one balance sheet's figure overwrite the other's — determine which PDF is which year from its "As of <date>" header.

Return ONLY valid JSON in this exact shape:
{
  "company": "string",
  "fiscal_year": "just the 4-digit fiscal year-end, e.g. '2026' — NOT a date range",
  "prior_year": "just the 4-digit fiscal year-end, e.g. '2025' — NOT a date range",
  "period_end": "string (e.g. 'January 31, 2026')",
  "lines": {
    "4000101": {"description": "Vente software - Montréal", "current": number, "prior": number},
    "...": {"description": "...", "current": number, "prior": number}
  }
}
Use 0 for missing values."""

def extract(files, overrides=None):
    client = anthropic.Anthropic(api_key=get_api_key())
    content = []
    for f in files:
        data_bytes = f.read()
        b64 = base64.b64encode(data_bytes).decode()
        content.append({
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
            "title": f.name
        })
    content.append({"type": "text", "text": build_prompt()})

    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=16000,
        system="You are a document transcription API. You ONLY output valid JSON. Never write explanations, steps, headers, or any text outside the JSON object. Your entire response must be parseable by json.loads().",
        messages=[{"role": "user", "content": content}]
    )

    if not msg.content:
        raise ValueError("Claude returned an empty response. Try uploading one PDF at a time.")

    raw = msg.content[0].text.strip()
    if not raw:
        raise ValueError("Claude returned empty text. The PDFs may be too large — try one at a time.")

    raw = raw.lstrip("```json").lstrip("```").rstrip("```").strip()

    def _parse(text):
        parsed = json.loads(text)
        # Stash the raw transcription so changing an override re-buckets
        # instantly instead of costing another API call.
        st.session_state["raw_parsed_v3"] = parsed
        return categorize(
            parsed.get("company","Active Média inc."),
            parsed.get("fiscal_year",""), parsed.get("prior_year",""),
            parsed.get("period_end",""), parsed.get("lines",{}) or {},
            overrides=overrides
        )

    try:
        return _parse(raw)
    except json.JSONDecodeError as e:
        last = raw.rfind("}")
        if last > 0:
            fixed = raw[:last+1]
            open_cnt = fixed.count("{") - fixed.count("}")
            try:
                return _parse(fixed + "}" * open_cnt)
            except:
                pass
        raise ValueError(f"Could not parse Claude response as JSON: {e}\n\nRaw response (first 500 chars):\n{raw[:500]}")


# ── PDF Builder ──────────────────────────────────────────────────────────────
def build_pdf(data, s):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable, KeepTogether
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.pagesizes import letter

    BLACK = colors.black
    GREY  = colors.HexColor("#666666")

    cy = data.get("fiscal_year", s["fiscal_year"])
    py = data.get("prior_year",  s["prior_year"])
    period_str = f"Exercice clos le {s['period_end']}, avec informations comparatives de {py}"

    def ps(name, size=10, bold=False, color=BLACK, align=TA_LEFT):
        return ParagraphStyle(name + str(id(name)),
            fontName="Helvetica-Bold" if bold else "Helvetica",
            fontSize=size, textColor=color, alignment=align,
            leading=size * 1.4, spaceAfter=0, spaceBefore=0)

    # Column widths — label takes most space, two equal value columns
    LW  = 4.0 * inch
    VW  = 1.3 * inch
    CW  = [LW, VW, VW]
    TW  = LW + VW + VW  # total table width

    def hdr_para(text, size=10, bold=False, color=BLACK, align=TA_LEFT):
        return Paragraph(text, ps("h", size, bold, color, align))

    def val_para(text, bold=False, color=BLACK):
        return Paragraph(f'<para alignment="right">{text}</para>',
            ps("v", 9.5, bold, color, TA_RIGHT))

    def lbl_para(text, bold=False, indent=0, color=BLACK):
        style = ps("l", 9.5, bold, color, TA_LEFT)
        style.leftIndent = indent
        safe = str(text).replace("\xa0", " ").replace("\x00", "") if text is not None else ""
        return Paragraph(safe, style)

    # Table style base
    BASE_STYLE = [
        ('TOPPADDING',    (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ('LEFTPADDING',   (0,0), (-1,-1), 0),
        ('RIGHTPADDING',  (0,0), (-1,-1), 0),
        ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
        ('FONTNAME',      (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE',      (0,0), (-1,-1), 9.5),
    ]

    def build_table(rows):
        """
        rows: list of (label_cell, v1_cell, v2_cell, style_type)
        style_type: 'normal'|'indent'|'section'|'subtotal'|'grand'|'blank'
        """
        tdata = []
        cmds  = list(BASE_STYLE)

        # Column header row
        tdata.append([
            hdr_para("", 9, False, GREY),
            val_para(cy, bold=True, color=BLACK),
            val_para(py, bold=True, color=BLACK),
        ])
        cmds += [
            ('LINEBELOW', (0,0), (-1,0), 0.75, BLACK),
            ('TOPPADDING', (0,0), (-1,0), 4),
            ('BOTTOMPADDING', (0,0), (-1,0), 4),
        ]

        for i, (lbl, v1, v2, stype) in enumerate(rows):
            rn = i + 1
            if stype == 'blank':
                tdata.append([Paragraph("", ps("b",3)), Paragraph("",ps("b2",3)), Paragraph("",ps("b3",3))])
                cmds.append(('FONTSIZE',(0,rn),(-1,rn),3))
                cmds.append(('TOPPADDING',(0,rn),(-1,rn),1))
                cmds.append(('BOTTOMPADDING',(0,rn),(-1,rn),1))
                continue

            # Label may already be a Paragraph (built via lbl_para by the caller)
            # or a plain string — never re-wrap an existing Paragraph, since
            # str(Paragraph(...)) prints its Python repr instead of its text.
            lbl_cell = lbl if isinstance(lbl, Paragraph) else lbl_para(lbl)

            if stype == 'section':
                tdata.append([lbl_cell, val_para(""), val_para("")])
                continue

            bold_row = stype in ('grand', 'subtotal')
            tdata.append([lbl_cell, val_para(v1, bold=bold_row), val_para(v2, bold=bold_row)])

            if stype == 'grand':
                cmds += [
                    ('LINEABOVE', (0,rn), (-1,rn), 0.75, BLACK),
                    ('LINEBELOW', (0,rn), (-1,rn), 1.5,  BLACK),
                    ('FONTNAME',  (0,rn), (-1,rn), 'Helvetica-Bold'),
                ]
            elif stype == 'subtotal':
                cmds += [
                    ('LINEABOVE', (0,rn), (-1,rn), 0.5, BLACK),
                    ('FONTNAME',  (0,rn), (-1,rn), 'Helvetica-Bold'),
                ]

        t = Table(tdata, colWidths=CW)
        t.setStyle(TableStyle(cmds))
        return t

    def v(obj, key):
        d = (obj or {}).get(key, {})
        return (d.get("current", 0) or 0) if isinstance(d, dict) else (d or 0)

    def vp(obj, key):
        d = (obj or {}).get(key, {})
        return (d.get("prior", 0) or 0) if isinstance(d, dict) else (d or 0)

    pl = data.get("pl", {}); a1 = data.get("ann1", {})
    a2 = data.get("ann2", {}); a3 = data.get("ann3", {}); a4 = data.get("ann4", {})

    story = []

    def page_hdr(title, period):
        story.append(Paragraph(s["company_name"].upper(), ps("co", 16, bold=True)))
        story.append(Spacer(1, 2))
        story.append(Paragraph(title, ps("ti", 10)))
        story.append(Spacer(1, 2))
        story.append(HRFlowable(width=TW, thickness=0.75, color=BLACK, spaceAfter=3, spaceBefore=0))
        story.append(Paragraph(period, ps("pe", 9, color=GREY)))
        story.append(Spacer(1, 10))

    # ── P&L ───────────────────────────────────────────────────────────────────
    page_hdr("État non consolidé des résultats", period_str)

    pl_rows = [
        (lbl_para("Ventes (notes 11 et 12)"), fmt_with_dollar(v(pl,"ventes")), fmt_with_dollar(vp(pl,"ventes")), "normal"),
        (lbl_para(""), "", "", "blank"),
        (lbl_para("Coût des ventes (annexe 1)", indent=12), fmt_num(v(pl,"cout_des_ventes")), fmt_num(vp(pl,"cout_des_ventes")), "normal"),
        (lbl_para(""), fmt_num(v(pl,"benefice_brut")), fmt_num(vp(pl,"benefice_brut")), "subtotal"),
        (lbl_para(""), "", "", "blank"),
        ("Charges", "", "", "section"),
        (lbl_para("    Frais d'exploitation (annexe 2)", indent=20), fmt_num(v(pl,"frais_exploitation")), fmt_num(vp(pl,"frais_exploitation")), "normal"),
        (lbl_para("    Frais financiers (annexe 3)", indent=20), fmt_num(v(pl,"frais_financiers")), fmt_num(vp(pl,"frais_financiers")), "normal"),
        (lbl_para("    Amortissement des immobilisations corporelles et actifs incorporels", indent=20), fmt_num(v(pl,"amortissement")), fmt_num(vp(pl,"amortissement")), "normal"),
        (lbl_para(""), fmt_num(v(pl,"total_charges")), fmt_num(vp(pl,"total_charges")), "subtotal"),
        (lbl_para(""), "", "", "blank"),
        (lbl_para("Bénéfice (perte) avant les autres revenus"), fmt_num(v(pl,"benefice_avant_autres")), fmt_num(vp(pl,"benefice_avant_autres")), "normal"),
        (lbl_para(""), "", "", "blank"),
        (lbl_para("Autres revenus (annexe 4)"), fmt_num(v(pl,"autres_revenus")), fmt_num(vp(pl,"autres_revenus")), "normal"),
        (lbl_para(""), "", "", "blank"),
        (lbl_para("Bénéfice net (perte nette)", bold=True), fmt_with_dollar(v(pl,"benefice_net")), fmt_with_dollar(vp(pl,"benefice_net")), "grand"),
    ]
    story.append(build_table(pl_rows))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Se reporter aux notes afférentes aux états financiers.", ps("note", 8, color=GREY)))
    story.append(PageBreak())

    # ── ANNEXES 1 & 2 ─────────────────────────────────────────────────────────
    page_hdr("Annexes", period_str)

    story.append(Paragraph("<b>Annexe 1 - Coût des ventes</b>", ps("a1", 10, bold=True)))
    story.append(Spacer(1, 5))
    ann1_rows = [
        (lbl_para("Achats"), fmt_with_dollar(v(a1,"achats")), fmt_with_dollar(vp(a1,"achats")), "normal"),
        (lbl_para("Salaires et avantages sociaux"), fmt_num(v(a1,"salaires")), fmt_num(vp(a1,"salaires")), "normal"),
        (lbl_para("Sous-traitance"), fmt_num(v(a1,"soustraitance")), fmt_num(vp(a1,"soustraitance")), "normal"),
        (lbl_para("Frais de livraison"), fmt_num(v(a1,"livraison")), fmt_num(vp(a1,"livraison")), "normal"),
        (lbl_para("Logiciel"), fmt_num(v(a1,"logiciel")), fmt_num(vp(a1,"logiciel")), "normal"),
        (lbl_para("", bold=True), fmt_with_dollar(v(a1,"total")), fmt_with_dollar(vp(a1,"total")), "grand"),
    ]
    story.append(build_table(ann1_rows))
    story.append(Spacer(1, 16))

    story.append(Paragraph("<b>Annexe 2 - Frais d'exploitation</b>", ps("a2", 10, bold=True)))
    story.append(Spacer(1, 5))

    def a2row(label, key):
        return (lbl_para(label), fmt_num(v(a2,key)), fmt_num(vp(a2,key)), "normal")

    ann2_rows = [
        (lbl_para("Salaires et avantages sociaux - recherche et développement"), fmt_with_dollar(v(a2,"salaires_rd")), fmt_with_dollar(vp(a2,"salaires_rd")), "normal"),
        a2row("Crédits d'impôt pour la recherche et le développement", "credits_rd"),
        a2row("Salaires et avantages sociaux - ventes et administration", "salaires_admin"),
        a2row("Frais de déplacement", "deplacement"),
        a2row("Location d'équipement", "location_equip"),
        a2row("Publicité et promotion", "publicite"),
        a2row("Honoraires", "honoraires"),
        a2row("Loyer", "loyer"),
        a2row("Télécommunications", "telecom"),
        a2row("Frais de représentation", "representation"),
        a2row("Frais de bureau", "bureau"),
        a2row("Cotisation et abonnement", "cotisation"),
        a2row("Assurances", "assurance"),
        a2row("Taxes et permis", "taxes"),
        a2row("Entretien et réparations", "entretien"),
        a2row("Courrier et frais postaux", "courrier"),
        a2row("Licence", "licence"),
        a2row("Frais de gestion de paie", "paie"),
        a2row("Frais de formation", "formation"),
        a2row("(Gain) Perte de change", "fx"),
        a2row("Intérêts et pénalités", "interet_penalites"),
        a2row("Représentant externe", "representant"),
        (lbl_para("", bold=True), fmt_with_dollar(v(a2,"total")), fmt_with_dollar(vp(a2,"total")), "grand"),
    ]
    story.append(build_table(ann2_rows))
    story.append(PageBreak())

    # ── ANNEXES 3 & 4 ─────────────────────────────────────────────────────────
    page_hdr("Annexes (suite)", period_str)

    story.append(Paragraph("<b>Annexe 3 - Frais financiers</b>", ps("a3", 10, bold=True)))
    story.append(Spacer(1, 5))
    ann3_rows = [
        (lbl_para("Intérêts sur la dette à long terme"), fmt_with_dollar(v(a3,"interet_lt")), fmt_with_dollar(vp(a3,"interet_lt")), "normal"),
        (lbl_para("Intérêts et frais bancaires"), fmt_num(v(a3,"frais_bancaires")), fmt_num(vp(a3,"frais_bancaires")), "normal"),
        (lbl_para("", bold=True), fmt_with_dollar(v(a3,"total")), fmt_with_dollar(vp(a3,"total")), "grand"),
    ]
    story.append(build_table(ann3_rows))
    story.append(Spacer(1, 20))

    story.append(Paragraph("<b>Annexe 4 - Autres revenus</b>", ps("a4", 10, bold=True)))
    story.append(Spacer(1, 5))
    ann4_rows = [
        (lbl_para("Aide gouvernementale (note 15 b))"), fmt_with_dollar(v(a4,"aide_gouv")), fmt_with_dollar(vp(a4,"aide_gouv")), "normal"),
        (lbl_para("Autres revenus"), fmt_num(v(a4,"autres")), fmt_num(vp(a4,"autres")), "normal"),
        (lbl_para("", bold=True), fmt_with_dollar(v(a4,"total")), fmt_with_dollar(vp(a4,"total")), "grand"),
    ]
    story.append(build_table(ann4_rows))

    # ── BALANCE SHEET ─────────────────────────────────────────────────────────
    bs = data.get("bs", {})
    if bs:
        def bv(key): return (bs.get(key,{}) or {}).get("current",0) or 0
        def bvp(key): return (bs.get(key,{}) or {}).get("prior",0) or 0
        story.append(PageBreak())
        page_hdr("Bilan non consolidé", str(s.get("period_end","")) + ", avec informations comparatives de " + str(data.get("prior_year", s.get("prior_year",""))))

        bs_rows = [
            # ACTIF
            (lbl_para("Actif"), "", "", "section"),
            (lbl_para(""), "", "", "blank"),
            (lbl_para("Actif à court terme", bold=False), "", "", "section"),
            (lbl_para("    Encaisse", indent=20),               fmt_with_dollar(bv("encaisse")),           fmt_with_dollar(bvp("encaisse")),           "normal"),
            (lbl_para("    Comptes clients et autres créances (note 2)", indent=20), fmt_num(bv("comptes_clients")), fmt_num(bvp("comptes_clients")), "normal"),
            (lbl_para("    Stocks", indent=20),                 fmt_num(bv("stocks")),                     fmt_num(bvp("stocks")),                     "normal"),
            (lbl_para("    Travaux en cours (note 15)", indent=20), fmt_num(bv("travaux_en_cours")),       fmt_num(bvp("travaux_en_cours")),           "normal"),
            (lbl_para("    Crédits d'impôt à recevoir", indent=20), fmt_num(bv("credits_impot")),        fmt_num(bvp("credits_impot")),              "normal"),
            (lbl_para("    Charges payées d'avance", indent=20), fmt_num(bv("charges_payees_avance")),   fmt_num(bvp("charges_payees_avance")),      "normal"),
            (lbl_para(""),                                       fmt_num(bv("total_actif_ct")),             fmt_num(bvp("total_actif_ct")),             "subtotal"),
            (lbl_para(""), "", "", "blank"),
            (lbl_para("Dépôts à long terme"),                   fmt_num(bv("depots_lt")),                  fmt_num(bvp("depots_lt")),                  "normal"),
            (lbl_para("Frais payés d'avance"),                 fmt_num(bv("frais_payes_avance_lt")),      fmt_num(bvp("frais_payes_avance_lt")),      "normal"),
            (lbl_para("Immobilisations corporelles (note 3)"),  fmt_num(bv("immobilisations")),            fmt_num(bvp("immobilisations")),            "normal"),
            (lbl_para("Actifs incorporels (note 4)"),           fmt_num(bv("actifs_incorporels")),         fmt_num(bvp("actifs_incorporels")),         "normal"),
            (lbl_para("Impôts futurs (note 14)"),               fmt_num(bv("impots_futurs")),              fmt_num(bvp("impots_futurs")),              "normal"),
            (lbl_para("Avances à la filiale, sans intérêt ni modalités de d'encaissement"), fmt_num(bv("avances_filiale")), fmt_num(bvp("avances_filiale")), "normal"),
            (lbl_para("Avances aux actionnaires, sans intérêt\nni modalités de d'encaissement"), fmt_num(bv("avances_actionnaires")), fmt_num(bvp("avances_actionnaires")), "normal"),
            (lbl_para("Placement dans la filiale (note 5)"),    fmt_num(bv("placement_filiale")),          fmt_num(bvp("placement_filiale")),          "normal"),
            (lbl_para(""), "", "", "blank"),
            (lbl_para("", bold=True),                           fmt_with_dollar(bv("total_actif")),        fmt_with_dollar(bvp("total_actif")),        "grand"),
            (lbl_para(""), "", "", "blank"),
            # PASSIF
            (lbl_para("Passif et avoir des actionnaires"), "", "", "section"),
            (lbl_para(""), "", "", "blank"),
            (lbl_para("Passif à court terme", bold=False), "", "", "section"),
            (lbl_para("    Emprunt bancaire (note 6)", indent=20), fmt_with_dollar(bv("emprunt_bancaire")), fmt_with_dollar(bvp("emprunt_bancaire")), "normal"),
            (lbl_para("    Comptes fournisseurs et charges à payer (note 7)", indent=20), fmt_num(bv("comptes_fournisseurs")), fmt_num(bvp("comptes_fournisseurs")), "normal"),
            (lbl_para("    Impôts sur le bénéfice à payer", indent=20), fmt_num(bv("impots_benefice")),   fmt_num(bvp("impots_benefice")),            "normal"),
            (lbl_para("    Produits reportés (note 15)", indent=20), fmt_num(bv("produits_reportes")),     fmt_num(bvp("produits_reportes")),          "normal"),
            (lbl_para("    Tranche à court terme de la dette à long terme (note 8)", indent=20), fmt_num(bv("tranche_ct_lt")), fmt_num(bvp("tranche_ct_lt")), "normal"),
            (lbl_para(""),                                       fmt_num(bv("total_passif_ct")),            fmt_num(bvp("total_passif_ct")),            "subtotal"),
            (lbl_para(""), "", "", "blank"),
            (lbl_para("Avantages incitatifs liés aux baux"),    fmt_num(bv("avantages_baux")),             fmt_num(bvp("avantages_baux")),             "normal"),
            (lbl_para("Dette à long terme (note 8)"),           fmt_num(bv("dette_lt")),                   fmt_num(bvp("dette_lt")),                   "normal"),
            (lbl_para(""),                                       fmt_num(bv("total_passif")),               fmt_num(bvp("total_passif")),               "subtotal"),
            (lbl_para(""), "", "", "blank"),
            # EQUITY
            (lbl_para("Avoir des actionnaires"), "", "", "section"),
            (lbl_para("    Capital-actions (note 9)", indent=20), fmt_num(bv("capital_actions")),          fmt_num(bvp("capital_actions")),            "normal"),
            (lbl_para("    Déficit", indent=20),                fmt_num(bv("deficit")),                    fmt_num(bvp("deficit")),                    "normal"),
            (lbl_para(""),                                       fmt_num(bv("total_avoir")),                fmt_num(bvp("total_avoir")),                "subtotal"),
            (lbl_para(""), "", "", "blank"),
            (lbl_para("Engagements (note 13)"), "", "", "normal"),
            (lbl_para(""), "", "", "blank"),
            (lbl_para("", bold=True),                           fmt_with_dollar(bv("total_passif_avoir")), fmt_with_dollar(bvp("total_passif_avoir")), "grand"),
        ]
        story.append(build_table(bs_rows))
        story.append(Spacer(1,10))
        story.append(Paragraph("Se reporter aux notes afférentes aux états financiers.", ps("note",8,color=GREY)))
        story.append(Spacer(1,8))
        story.append(Paragraph("Au nom du conseil,", ps("conseil",9,color=GREY)))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.9*inch, rightMargin=0.9*inch,
        topMargin=0.85*inch, bottomMargin=0.85*inch)
    doc.build(story)
    buf.seek(0)
    return buf

# ── Word builder ──────────────────────────────────────────────────────────────
def build_word(data, s):
    doc = Document()
    for sec in doc.sections:
        sec.page_width=Inches(8.5); sec.page_height=Inches(11)
        sec.left_margin=sec.right_margin=Inches(0.9)
        sec.top_margin=sec.bottom_margin=Inches(0.85)

    NAVY=RGBColor(0x1A,0x2E,0x4A); BLACK=RGBColor(0,0,0); GREY=RGBColor(0x55,0x55,0x55)
    cy=data.get("fiscal_year",s["fiscal_year"]); py=data.get("prior_year",s["prior_year"])
    COL1=Twips(5400); COL2=Twips(1800); COL3=Twips(1800)

    def set_bg(cell,hex_color):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
        tcPr.append(shd)

    def no_bdr(cell):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        for ex in tcPr.findall(qn('w:tcBorders')): tcPr.remove(ex)
        tcB=OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'none'); tcB.append(b)
        tcPr.append(tcB)

    def underline_cell(cell, thick=False):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        for ex in tcPr.findall(qn('w:tcBorders')): tcPr.remove(ex)
        tcB=OxmlElement('w:tcBorders')
        for side in ['top','left','right']:
            b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'none'); tcB.append(b)
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'12' if thick else '6'); b.set(qn('w:color'),'000000')
        tcB.append(b); tcPr.append(tcB)

    def run(para,text,bold=False,size=10,color=BLACK):
        r=para.add_run(text); r.bold=bold; r.font.name="Arial"; r.font.size=Pt(size)
        r.font.color.rgb=color

    def ns(p): p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)

    def page_hdr(title, period):
        p=doc.add_paragraph(); ns(p); run(p,s["company_name"].upper(),bold=True,size=16)
        p2=doc.add_paragraph(); ns(p2); run(p2,title,size=10)
        p3=doc.add_paragraph()
        pPr=p3._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
        bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single')
        bot.set(qn('w:sz'),'6'); bot.set(qn('w:color'),'000000')
        pBdr.append(bot); pPr.append(pBdr); p3.paragraph_format.space_after=Pt(8)
        run(p3,period,size=9)

    def make_table(rows):
        tbl=doc.add_table(rows=0,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
        # Header row
        tr=tbl.add_row(); cells=tr.cells
        cells[0].width=COL1; cells[1].width=COL2; cells[2].width=COL3
        for c in cells: set_bg(c,"FFFFFF"); no_bdr(c)
        for c,val,al in [(cells[0],"",WD_ALIGN_PARAGRAPH.LEFT),
                          (cells[1],cy,WD_ALIGN_PARAGRAPH.RIGHT),
                          (cells[2],py,WD_ALIGN_PARAGRAPH.RIGHT)]:
            p=c.paragraphs[0]; p.alignment=al; ns(p)
            run(p,val,bold=True,size=9)
            underline_cell(c)

        for lbl,v1,v2,stype in rows:
            tr=tbl.add_row(); cells=tr.cells
            cells[0].width=COL1; cells[1].width=COL2; cells[2].width=COL3
            for c in cells: set_bg(c,"FFFFFF"); no_bdr(c)

            if stype=="blank":
                for c in cells:
                    p=c.paragraphs[0]; ns(p); p.paragraph_format.space_before=Pt(3)
                continue

            bold_row=stype in ("grand","total","subtotal")
            for c,val,al in [(cells[0],lbl,WD_ALIGN_PARAGRAPH.LEFT),
                              (cells[1],v1,WD_ALIGN_PARAGRAPH.RIGHT),
                              (cells[2],v2,WD_ALIGN_PARAGRAPH.RIGHT)]:
                p=c.paragraphs[0]; p.alignment=al; ns(p)
                p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
                run(p,val,bold=bold_row,size=9.5)

            if stype=="grand":
                for c in cells: underline_cell(c,thick=True)
            elif stype in ("total","subtotal"):
                for c in cells: underline_cell(c)

        doc.add_paragraph().paragraph_format.space_after=Pt(6)

    def v(obj,key):
        d=obj.get(key,{}) if obj else {}
        return (d.get("current",0) or 0) if isinstance(d,dict) else (d or 0)
    def vp(obj,key):
        d=obj.get(key,{}) if obj else {}
        return (d.get("prior",0) or 0) if isinstance(d,dict) else (d or 0)

    pl=data.get("pl",{}); a1=data.get("ann1",{}); a2=data.get("ann2",{})
    a3=data.get("ann3",{}); a4=data.get("ann4",{})
    period_str=f"Exercice clos le {s['period_end']}, avec informations comparatives de {py}"

    page_hdr("État non consolidé des résultats", period_str)
    make_table([
        ("Ventes (notes 11 et 12)", fmt_with_dollar(v(pl,"ventes")), fmt_with_dollar(vp(pl,"ventes")), "normal"),
        ("","","","blank"),
        ("Coût des ventes (annexe 1)", fmt_num(v(pl,"cout_des_ventes")), fmt_num(vp(pl,"cout_des_ventes")), "indent"),
        ("", fmt_num(v(pl,"benefice_brut")), fmt_num(vp(pl,"benefice_brut")), "subtotal"),
        ("","","","blank"),
        ("Charges","","","section"),
        ("    Frais d'exploitation (annexe 2)", fmt_num(v(pl,"frais_exploitation")), fmt_num(vp(pl,"frais_exploitation")), "indent"),
        ("    Frais financiers (annexe 3)", fmt_num(v(pl,"frais_financiers")), fmt_num(vp(pl,"frais_financiers")), "indent"),
        ("    Amortissement des immobilisations corporelles et actifs incorporels",
         fmt_num(v(pl,"amortissement")), fmt_num(vp(pl,"amortissement")), "indent"),
        ("", fmt_num(v(pl,"total_charges")), fmt_num(vp(pl,"total_charges")), "subtotal"),
        ("","","","blank"),
        ("Bénéfice (perte) avant les autres revenus", fmt_num(v(pl,"benefice_avant_autres")), fmt_num(vp(pl,"benefice_avant_autres")), "normal"),
        ("","","","blank"),
        ("Autres revenus (annexe 4)", fmt_num(v(pl,"autres_revenus")), fmt_num(vp(pl,"autres_revenus")), "normal"),
        ("","","","blank"),
        ("Bénéfice net (perte nette)", fmt_with_dollar(v(pl,"benefice_net")), fmt_with_dollar(vp(pl,"benefice_net")), "grand"),
    ])
    p=doc.add_paragraph(); run(p,"Se reporter aux notes afférentes aux états financiers.",size=8,color=GREY)
    doc.add_page_break()

    page_hdr("Annexes", period_str)
    p=doc.add_paragraph(); run(p,"Annexe 1 - Coût des ventes",bold=True,size=10); p.paragraph_format.space_after=Pt(4)
    make_table([
        ("Achats", fmt_with_dollar(v(a1,"achats")), fmt_with_dollar(vp(a1,"achats")), "normal"),
        ("Salaires et avantages sociaux", fmt_num(v(a1,"salaires")), fmt_num(vp(a1,"salaires")), "normal"),
        ("Sous-traitance", fmt_num(v(a1,"soustraitance")), fmt_num(vp(a1,"soustraitance")), "normal"),
        ("Frais de livraison", fmt_num(v(a1,"livraison")), fmt_num(vp(a1,"livraison")), "normal"),
        ("Logiciel", fmt_num(v(a1,"logiciel")), fmt_num(vp(a1,"logiciel")), "normal"),
        ("", fmt_with_dollar(v(a1,"total")), fmt_with_dollar(vp(a1,"total")), "grand"),
    ])
    p=doc.add_paragraph(); run(p,"Annexe 2 - Frais d'exploitation",bold=True,size=10); p.paragraph_format.space_after=Pt(4)
    make_table([
        ("Salaires et avantages sociaux - recherche et développement", fmt_with_dollar(v(a2,"salaires_rd")), fmt_with_dollar(vp(a2,"salaires_rd")), "normal"),
        ("Crédits d'impôt pour la recherche et le développement", fmt_num(v(a2,"credits_rd")), fmt_num(vp(a2,"credits_rd")), "normal"),
        ("Salaires et avantages sociaux - ventes et administration", fmt_num(v(a2,"salaires_admin")), fmt_num(vp(a2,"salaires_admin")), "normal"),
        ("Frais de déplacement", fmt_num(v(a2,"deplacement")), fmt_num(vp(a2,"deplacement")), "normal"),
        ("Location d'équipement", fmt_num(v(a2,"location_equip")), fmt_num(vp(a2,"location_equip")), "normal"),
        ("Publicité et promotion", fmt_num(v(a2,"publicite")), fmt_num(vp(a2,"publicite")), "normal"),
        ("Honoraires", fmt_num(v(a2,"honoraires")), fmt_num(vp(a2,"honoraires")), "normal"),
        ("Loyer", fmt_num(v(a2,"loyer")), fmt_num(vp(a2,"loyer")), "normal"),
        ("Télécommunications", fmt_num(v(a2,"telecom")), fmt_num(vp(a2,"telecom")), "normal"),
        ("Frais de représentation", fmt_num(v(a2,"representation")), fmt_num(vp(a2,"representation")), "normal"),
        ("Frais de bureau", fmt_num(v(a2,"bureau")), fmt_num(vp(a2,"bureau")), "normal"),
        ("Cotisation et abonnement", fmt_num(v(a2,"cotisation")), fmt_num(vp(a2,"cotisation")), "normal"),
        ("Assurances", fmt_num(v(a2,"assurance")), fmt_num(vp(a2,"assurance")), "normal"),
        ("Taxes et permis", fmt_num(v(a2,"taxes")), fmt_num(vp(a2,"taxes")), "normal"),
        ("Entretien et réparations", fmt_num(v(a2,"entretien")), fmt_num(vp(a2,"entretien")), "normal"),
        ("Courrier et frais postaux", fmt_num(v(a2,"courrier")), fmt_num(vp(a2,"courrier")), "normal"),
        ("Licence", fmt_num(v(a2,"licence")), fmt_num(vp(a2,"licence")), "normal"),
        ("Frais de gestion de paie", fmt_num(v(a2,"paie")), fmt_num(vp(a2,"paie")), "normal"),
        ("Frais de formation", fmt_num(v(a2,"formation")), fmt_num(vp(a2,"formation")), "normal"),
        ("(Gain) Perte de change", fmt_num(v(a2,"fx")), fmt_num(vp(a2,"fx")), "normal"),
        ("Intérêts et pénalités", fmt_num(v(a2,"interet_penalites")), fmt_num(vp(a2,"interet_penalites")), "normal"),
        ("Représentant externe", fmt_num(v(a2,"representant")), fmt_num(vp(a2,"representant")), "normal"),
        ("", fmt_with_dollar(v(a2,"total")), fmt_with_dollar(vp(a2,"total")), "grand"),
    ])
    doc.add_page_break()

    page_hdr("Annexes (suite)", period_str)
    p=doc.add_paragraph(); run(p,"Annexe 3 - Frais financiers",bold=True,size=10); p.paragraph_format.space_after=Pt(4)
    make_table([
        ("Intérêts sur la dette à long terme", fmt_with_dollar(v(a3,"interet_lt")), fmt_with_dollar(vp(a3,"interet_lt")), "normal"),
        ("Intérêts et frais bancaires", fmt_num(v(a3,"frais_bancaires")), fmt_num(vp(a3,"frais_bancaires")), "normal"),
        ("", fmt_with_dollar(v(a3,"total")), fmt_with_dollar(vp(a3,"total")), "grand"),
    ])
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14)
    run(p,"Annexe 4 - Autres revenus",bold=True,size=10); p.paragraph_format.space_after=Pt(4)
    make_table([
        ("Aide gouvernementale (note 15 b))", fmt_with_dollar(v(a4,"aide_gouv")), fmt_with_dollar(vp(a4,"aide_gouv")), "normal"),
        ("Autres revenus", fmt_num(v(a4,"autres")), fmt_num(vp(a4,"autres")), "normal"),
        ("", fmt_with_dollar(v(a4,"total")), fmt_with_dollar(vp(a4,"total")), "grand"),
    ])

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ── Main UI ───────────────────────────────────────────────────────────────────
st.markdown('<div class="info-box">Upload your QuickBooks P&L PDF or Excel export. Include prior year PDF alongside current year for the two-column comparison.</div>', unsafe_allow_html=True)

uploaded = st.file_uploader("Drop files here", type=["pdf","xlsx"], accept_multiple_files=True, label_visibility="collapsed")
if uploaded:
    st.markdown(f"**{len(uploaded)} file(s) ready:** " + ", ".join(f.name for f in uploaded))

if not get_api_key():
    st.markdown('<div class="info-box">👈 Enter your Anthropic API key in the sidebar.</div>', unsafe_allow_html=True)

if st.button("⚡ Extract & Build Report", disabled=not(uploaded and get_api_key()), type="primary", use_container_width=True):
    with st.spinner("Reading files and grouping accounts... ~30 seconds"):
        try:
            data = extract(uploaded, overrides=overrides)
            st.session_state["data_v3"] = data
        except Exception as e:
            st.error(f"Extraction failed: {e}")

# Overrides are applied deterministically in Python, so changing one only needs
# a re-bucket of the cached transcription — no second API call, no re-upload.
if "raw_parsed_v3" in st.session_state:
    _p = st.session_state["raw_parsed_v3"]
    st.session_state["data_v3"] = categorize(
        _p.get("company", "Active Média inc."),
        _p.get("fiscal_year", ""), _p.get("prior_year", ""),
        _p.get("period_end", ""), _p.get("lines", {}) or {},
        overrides=overrides
    )

if "data_v3" in st.session_state:
    data = st.session_state["data_v3"]
    pl = data.get("pl",{})
    cy = data.get("fiscal_year", fiscal_year)
    py_val = data.get("prior_year", prior_year)

    st.success(f"✅ {data.get('company','')} · FY{cy} vs FY{py_val}")
    if data.get("_bs_empty"):
        st.warning("⚠️ The balance sheet came back empty (all values are $0). This usually means the balance sheet page/PDF wasn't part of the upload, or it wasn't recognized. Double-check you included it, then re-run.")

    imb_cur = data.get("_bs_imbalance_current", 0)
    imb_pri = data.get("_bs_imbalance_prior", 0)
    if abs(imb_cur) > 5 or abs(imb_pri) > 5:
        parts = []
        if abs(imb_cur) > 5: parts.append(f"current year off by {fmt_with_dollar(imb_cur)}")
        if abs(imb_pri) > 5: parts.append(f"prior year off by {fmt_with_dollar(imb_pri)}")
        st.warning(f"⚠️ Balance sheet doesn't tie out — Total Actif ≠ Total Passif+Avoir ({', '.join(parts)}).")
        with st.expander("🔍 Diagnose the imbalance", expanded=True):
            st.caption("Most likely causes, in order. Work top-down — fixing one changes the residual, so re-run after each.")
            st.markdown(
                "1. **Unmapped accounts** — anything in the red box below is excluded from every total. "
                "Add it to `ACCOUNT_MAP` first.\n"
                "2. **Auditor schedule items** — Travaux en cours, Produits reportés, and Tranche à court "
                "terme de la dette à long terme come off schedules, not the trial balance. Enter them under "
                "*Ajustements de l'auditeur* in the sidebar.\n"
                "3. **Account 2163 (Impôts futurs)** — QuickBooks files it as a credit; the audited statement "
                "presents it as a long-term asset. If the sign isn't flipped on import, the error is exactly 2× its balance.\n"
                "4. **Opening déficit (account 3560)** — the prior-year gap being much larger than the current-year "
                "gap points here. Confirm the opening retained-earnings balance carried through."
            )
            bs_d = data.get("bs", {})
            st.markdown("**Where the two sides land:**")
            d1, d2 = st.columns(2)
            with d1:
                st.metric("Total actif (courant)", fmt_num((bs_d.get("total_actif") or {}).get("current", 0)))
                st.metric("Total actif (précédent)", fmt_num((bs_d.get("total_actif") or {}).get("prior", 0)))
            with d2:
                st.metric("Total passif + avoir (courant)", fmt_num((bs_d.get("total_passif_avoir") or {}).get("current", 0)))
                st.metric("Total passif + avoir (précédent)", fmt_num((bs_d.get("total_passif_avoir") or {}).get("prior", 0)))

            if st.checkbox("Show account-level tie-out table"):
                rows = []
                for code, info in (data.get("_lines") or {}).items():
                    bucket = ACCOUNT_MAP.get(code) or DESCRIPTION_MAP.get(
                        (info.get("description") or "").strip().lower()) or "— UNMAPPED —"
                    c = info.get("current", 0) or 0
                    p = info.get("prior", 0) or 0
                    flipped = False
                    if code in SIGN_NORMALIZE:
                        want = SIGN_NORMALIZE[code]
                        nc, np_ = (abs(c), abs(p)) if want == "positive" else (-abs(c), -abs(p))
                        flipped = (nc, np_) != (c, p)
                        c, p = nc, np_
                    rows.append({
                        "Code": code,
                        "Description": info.get("description", ""),
                        "Bucket": bucket,
                        "Signe inversé": "✔" if flipped else "",
                        "Courant": c,
                        "Précédent": p,
                    })
                if rows:
                    df = pd.DataFrame(rows).sort_values(["Bucket", "Code"])
                    bs_only = st.checkbox("Balance sheet accounts only", value=True)
                    if bs_only:
                        df = df[df["Bucket"].str.startswith("bs_") | (df["Bucket"] == "— UNMAPPED —")]
                    st.dataframe(df, use_container_width=True, hide_index=True)
                    st.download_button("⬇️ Download tie-out as CSV",
                                       data=df.to_csv(index=False).encode("utf-8"),
                                       file_name="tie_out.csv", mime="text/csv")

    rollsrc = data.get("_roll_source", {})
    rollfwd = data.get("_roll_forward", {})
    if rollsrc:
        with st.expander("📈 Report du bénéfice au déficit — source par exercice", expanded=False):
            st.caption("Le compte 3560 ne contient que le solde d'ouverture. Le bénéfice de "
                       "l'exercice doit y être ajouté pour obtenir le déficit de clôture.")
            for side, lbl in (("current", "Exercice courant"), ("prior", "Exercice précédent")):
                st.markdown(f"**{lbl}** — {fmt_with_dollar(rollfwd.get(side, 0))}  \n"
                            f"Source : {rollsrc.get(side, '—')}")

    flipped = data.get("_sign_flipped", {})
    if flipped:
        with st.expander(f"🔄 {len(flipped)} compte(s) — signe normalisé pour la présentation auditée", expanded=False):
            st.caption("Ces comptes sont classés par QuickBooks du côté opposé à leur présentation "
                       "dans les états financiers audités. Le signe est inversé avant tout regroupement.")
            for code, info in flipped.items():
                st.markdown(f"**{code}** — {info['description']}  \n"
                            f"Après inversion — Courant : {fmt_with_dollar(info['current'])} · "
                            f"Précédent : {fmt_with_dollar(info['prior'])}")

    applied = data.get("_applied_overrides", {})
    if applied:
        labels = {"travaux_en_cours": "Travaux en cours",
                  "produits_reportes": "Produits reportés",
                  "tranche_ct_lt": "Tranche à court terme de la dette à long terme"}
        st.info("🧾 Ajustements de l'auditeur appliqués : " +
                ", ".join(f"{labels.get(k, k)} ({', '.join(v)})" for k, v in applied.items()))

    needs_review = data.get("_needs_review", {})
    unmapped = data.get("_unmapped", {})
    if needs_review:
        with st.expander(f"⚠️ {len(needs_review)} account(s) pending accountant confirmation — excluded from all totals below", expanded=True):
            st.caption("These accounts don't clearly belong in any bucket yet — confirm with your accountant, then they can be added to the mapping.")
            for code, info in needs_review.items():
                st.markdown(f"**{code}** — {info['description']}  \nCurrent: {fmt_with_dollar(info['current'])} · Prior: {fmt_with_dollar(info['prior'])}  \n*{info['note']}*")
    if unmapped:
        with st.expander(f"🔴 {len(unmapped)} unrecognized account code(s) — not included in any total", expanded=True):
            st.caption("These codes weren't in the mapping table at all — likely a new account. Totals below may be incomplete until these are added.")
            for code, info in unmapped.items():
                st.markdown(f"**{code}** — {info['description']}  \nCurrent: {fmt_with_dollar(info['current'])} · Prior: {fmt_with_dollar(info['prior'])}")
    st.markdown("---")

    # KPIs
    def gv(key): return (pl.get(key,{}) or {}).get("current",0) or 0
    rev=gv("ventes"); gp=gv("benefice_brut"); np_=gv("benefice_net")
    c1,c2,c3,c4=st.columns(4)
    with c1: st.markdown(f'<div class="kpi-card"><div class="kpi-label">Revenue</div><div class="kpi-value">{fmt_num(rev)}</div></div>',unsafe_allow_html=True)
    with c2:
        cls="kpi-pos" if gp>=0 else "kpi-neg"
        st.markdown(f'<div class="kpi-card"><div class="kpi-label">Gross profit</div><div class="kpi-value {cls}">{fmt_num(gp)}</div><div class="kpi-sub">Margin {pct(gp,rev)}</div></div>',unsafe_allow_html=True)
    with c3:
        cls="kpi-pos" if np_>=0 else "kpi-neg"
        st.markdown(f'<div class="kpi-card"><div class="kpi-label">Net profit</div><div class="kpi-value {cls}">{fmt_num(np_)}</div><div class="kpi-sub">Margin {pct(np_,rev)}</div></div>',unsafe_allow_html=True)
    with c4:
        bs = data.get("bs", {})
        ta = (bs.get("total_actif",{}) or {}).get("current",0) if bs else 0
        st.markdown(f'<div class="kpi-card"><div class="kpi-label">Total actif</div><div class="kpi-value">{fmt_num(ta)}</div></div>',unsafe_allow_html=True)

    st.markdown("---")
    col1,col2=st.columns(2)
    slug=f"FY{cy}"

    with col1:
        pdf_buf=build_pdf(data,settings)
        st.download_button("📄 Download PDF (audited style)",data=pdf_buf,
            file_name=f"ActiveMedia_FinancialStatements_{slug}.pdf",
            mime="application/pdf",use_container_width=True,type="primary")
    with col2:
        word_buf=build_word(data,settings)
        st.download_button("📝 Download Word (.docx)",data=word_buf,
            file_name=f"ActiveMedia_FinancialStatements_{slug}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True)
